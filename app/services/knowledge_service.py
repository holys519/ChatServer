"""
Knowledge System Service
Simple implementation for document processing and knowledge management
"""

import os
import json
import uuid
from typing import List, Dict, Any, Optional
from datetime import datetime
from fastapi import UploadFile

from app.core.config import settings
from app.services.enhanced_document_processor import EnhancedDocumentProcessor
from app.models.enhanced_schemas import EnhancedChunk

# Firebase service import
try:
    from app.services.firebase_service import firebase_service
    FIREBASE_AVAILABLE = True
except ImportError:
    firebase_service = None
    FIREBASE_AVAILABLE = False

# Google Cloud Document AI import
try:
    from google.cloud import documentai
    DOCUMENT_AI_AVAILABLE = True
except ImportError:
    documentai = None
    DOCUMENT_AI_AVAILABLE = False

# Google Cloud Storage import
try:
    from google.cloud import storage
    STORAGE_AVAILABLE = True
except ImportError:
    storage = None
    STORAGE_AVAILABLE = False

# Vertex AI Embeddings import
try:
    from google.cloud import aiplatform
    from vertexai.language_models import TextEmbeddingModel
    VERTEX_AI_AVAILABLE = True
except ImportError:
    aiplatform = None
    TextEmbeddingModel = None
    VERTEX_AI_AVAILABLE = False

# Gemini service import for knowledge graph extraction
try:
    from app.services.gemini_service import gemini_service
    GEMINI_AVAILABLE = True
except ImportError:
    gemini_service = None
    GEMINI_AVAILABLE = False

# Enhanced knowledge graph service
try:
    from app.services.enhanced_knowledge_graph_service import enhanced_knowledge_graph_service
    ENHANCED_KG_AVAILABLE = True
except ImportError:
    enhanced_knowledge_graph_service = None
    ENHANCED_KG_AVAILABLE = False

class KnowledgeService:
    def __init__(self):
        """Initialize the Knowledge Service with Google Cloud integration"""
        print("🚀 Initializing Knowledge Service...")
        
        # Initialize Document AI client if available
        self.doc_ai_client = None
        if DOCUMENT_AI_AVAILABLE:
            try:
                self.doc_ai_client = documentai.DocumentProcessorServiceClient()
                print("✅ Document AI client initialized")
            except Exception as e:
                print(f"⚠️ Document AI initialization failed: {str(e)}")
        
        # Initialize Cloud Storage client if available
        self.storage_client = None
        if STORAGE_AVAILABLE:
            try:
                self.storage_client = storage.Client()
                print("✅ Cloud Storage client initialized")
            except Exception as e:
                print(f"⚠️ Cloud Storage initialization failed: {str(e)}")
        
        # Initialize Vertex AI if available
        self.embedding_model = None
        if VERTEX_AI_AVAILABLE:
            try:
                project_id = settings.google_cloud_project or settings.firebase_project_id
                location = settings.vertex_ai_location
                
                if project_id and location:
                    aiplatform.init(
                        project=project_id,
                        location=location
                    )
                    self.embedding_model = TextEmbeddingModel.from_pretrained("text-multilingual-embedding-002")
                    print("✅ Vertex AI Embeddings initialized")
                else:
                    print("⚠️ Vertex AI configuration missing (project_id or location)")
            except Exception as e:
                print(f"⚠️ Vertex AI initialization failed: {str(e)}")
        
        # Check Gemini service availability
        if GEMINI_AVAILABLE and gemini_service:
            print("✅ Gemini service available for knowledge graph extraction")
        else:
            print("⚠️ Gemini service not available")
        
        # Initialize enhanced document processor
        self.enhanced_processor = EnhancedDocumentProcessor(
            project_id=project_id,
            location=location
        )
        
        self.initialized = True
        print("✅ Knowledge service initialized with enhanced features")
    
    async def process_document(
        self, 
        job_id: str, 
        user_id: str, 
        file_content: bytes,
        file_name: str,
        content_type: str,
        options: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Process uploaded document with Firestore integration"""
        print(f"🚀 Starting document processing for job {job_id}")
        print(f"📄 Processing file: {file_name} (type: {content_type})")
        
        try:
            # File content is already available as bytes
            content = file_content
            file_size = len(content)
            
            print(f"📊 File size: {file_size} bytes")
            
            # Create job data
            job_data = {
                'id': job_id,
                'user_id': user_id,
                'document_name': file_name,
                'status': 'processing',
                'progress': 10,
                'current_step': 'Document received',
                'created_at': datetime.now(),
                'options': options,
                'file_size': file_size
            }
            
            # Save initial job to Firestore
            await self._save_processing_job(job_data)
            
            # Step 1: Upload to Cloud Storage
            await self._update_job_progress(job_id, 20, 'Uploading to Cloud Storage')
            storage_path = await self._upload_to_storage(job_id, user_id, content, file_name, content_type)
            
            # Step 2: Extract text with Document AI
            await self._update_job_progress(job_id, 30, 'Extracting text with Document AI')
            extracted_text = await self._extract_text_with_docai(storage_path, content_type, file_name)
            
            # Step 3: Chunk text for processing
            await self._update_job_progress(job_id, 50, 'Chunking text for embeddings')
            chunks = await self._chunk_text(extracted_text, options.get('chunk_size', 500), options.get('overlap_size', 50))
            
            # Step 4: Generate embeddings if enabled
            embeddings = None
            if options.get('enable_vector_search', True):
                await self._update_job_progress(job_id, 60, 'Generating vector embeddings')
                embeddings = await self._generate_embeddings(chunks)
                
                # Save vector chunks to Firestore
                await self._save_vector_chunks(job_id, user_id, chunks, embeddings)
            
            # Step 5: Extract knowledge graph if enabled
            entities = []
            relations = []
            if options.get('enable_knowledge_graph', True):
                await self._update_job_progress(job_id, 80, 'Extracting knowledge graph')
                knowledge_graph = await self._extract_knowledge_graph(extracted_text, chunks)
                entities = knowledge_graph.get('entities', [])
                relations = knowledge_graph.get('relations', [])
                
                # Save knowledge graph to Firestore
                await self._save_knowledge_graph(job_id, user_id, entities, relations)
            
            await self._update_job_progress(job_id, 100, 'Processing completed')
            
            # Create document record
            doc_data = {
                'id': job_id,  # Use job_id as document_id
                'user_id': user_id,
                'name': file.filename,
                'file_size': file_size,
                'chunk_count': len(chunks) if chunks else 0,
                'entity_count': len(entities),
                'processed_at': datetime.now(),
                'status': 'completed',
                'has_embeddings': embeddings is not None,
                'text_length': len(extracted_text)
            }
            
            await self._save_processed_document(doc_data)
            
            job_data.update({
                'status': 'completed',
                'progress': 100,
                'current_step': 'Processing completed',
                'completed_at': datetime.now()
            })
            
            await self._update_processing_job(job_id, job_data)
            
            print(f"✅ Document processing completed for job {job_id}")
            return job_data
            
        except Exception as e:
            print(f"❌ Document processing failed for job {job_id}: {str(e)}")
            error_data = {
                'id': job_id,
                'user_id': user_id,
                'document_name': file.filename,
                'status': 'failed',
                'progress': 0,
                'error_message': str(e),
                'created_at': datetime.now()
            }
            await self._save_processing_job(error_data)
            return error_data
    
    async def process_document_enhanced(
        self, 
        job_id: str, 
        user_id: str, 
        file_content: bytes,
        file_name: str,
        content_type: str,
        options: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Enhanced document processing with location tracking"""
        print(f"🚀 Starting enhanced document processing for job {job_id}")
        
        try:
            # Use enhanced processor for document analysis
            enhanced_chunks, processing_status = await self.enhanced_processor.process_document_with_location(
                file_content=file_content,
                file_name=file_name,
                content_type=content_type,
                job_id=job_id,
                chunk_size=options.get('chunk_size', 500),
                overlap_size=options.get('overlap_size', 50)
            )
            
            # Update processing status with user_id
            processing_status.user_id = user_id
            
            # Save initial job to Firestore
            await self._save_enhanced_processing_job(processing_status.dict())
            
            # Generate embeddings if enabled
            if options.get('enable_vector_search', True) and enhanced_chunks:
                await self._update_job_progress(job_id, 60, 'Generating vector embeddings')
                
                # Extract text for embedding generation
                chunk_texts = [chunk.text for chunk in enhanced_chunks]
                embeddings = await self._generate_embeddings(chunk_texts)
                
                if embeddings:
                    # Add embeddings to chunks
                    for chunk, embedding in zip(enhanced_chunks, embeddings):
                        chunk.embedding = embedding
                        chunk.embedding_model = "textembedding-gecko@003"
                
                # Save enhanced vector chunks to Firestore
                await self._save_enhanced_vector_chunks(job_id, user_id, enhanced_chunks)
            
            # Enhanced knowledge graph extraction (if enabled)
            if options.get('enable_knowledge_graph', True) and enhanced_chunks:
                await self._update_enhanced_job_progress(job_id, 80, 'knowledge_graph', 'Extracting enhanced knowledge graph')
                
                # Combine all chunk text for knowledge extraction
                full_text = "\n\n".join([chunk.text for chunk in enhanced_chunks])
                
                # Prepare chunk data for enhanced knowledge graph service
                chunk_data = []
                for chunk in enhanced_chunks:
                    chunk_data.append({
                        'id': chunk.id,
                        'text': chunk.text,
                        'metadata': chunk.metadata.dict()
                    })
                
                # Use enhanced knowledge graph service if available
                if ENHANCED_KG_AVAILABLE and enhanced_knowledge_graph_service:
                    try:
                        print("🧠 Using enhanced knowledge graph extraction")
                        knowledge_graph = await enhanced_knowledge_graph_service.extract_enhanced_knowledge_graph(
                            text=full_text,
                            chunks=chunk_data,
                            user_id=user_id,
                            document_id=job_id,
                            document_name=file_name
                        )
                        entities = knowledge_graph.get('entities', [])
                        relations = knowledge_graph.get('relations', [])
                        print(f"✅ Enhanced knowledge graph: {len(entities)} entities, {len(relations)} relations")
                    except Exception as e:
                        print(f"⚠️ Enhanced knowledge graph extraction failed: {str(e)}")
                        # Fallback to basic extraction
                        knowledge_graph = await self._extract_knowledge_graph(full_text, [])
                        entities = knowledge_graph.get('entities', [])
                        relations = knowledge_graph.get('relations', [])
                else:
                    print("⚠️ Enhanced knowledge graph service not available, using basic extraction")
                    knowledge_graph = await self._extract_knowledge_graph(full_text, [])
                    entities = knowledge_graph.get('entities', [])
                    relations = knowledge_graph.get('relations', [])
                
                # Save knowledge graph to Firestore (basic format for compatibility)
                await self._save_knowledge_graph(job_id, user_id, entities, relations)
            
            await self._update_job_progress(job_id, 100, 'Processing completed')
            
            # Create enhanced document record
            doc_data = {
                'id': job_id,
                'user_id': user_id,
                'name': file_name,
                'processing_status': 'completed',
                'file_size': len(file_content),
                'chunks_count': len(enhanced_chunks),
                'pages_count': processing_status.total_pages,
                'created_at': datetime.now(),
                'enhanced_processing': True,
                'processing_version': '2.0'
            }
            
            await self._save_processed_document(doc_data)
            
            return {
                'status': 'completed',
                'job_id': job_id,
                'chunks_count': len(enhanced_chunks),
                'pages_count': processing_status.total_pages,
                'enhanced': True
            }
            
        except Exception as e:
            print(f"❌ Enhanced document processing error: {str(e)}")
            
            error_data = {
                'id': job_id,
                'user_id': user_id,
                'document_name': file_name,
                'status': 'failed',
                'progress': 0,
                'error_message': str(e),
                'created_at': datetime.now(),
                'enhanced_processing': True
            }
            await self._save_enhanced_processing_job(error_data)
            return error_data
    
    async def _save_enhanced_processing_job(self, job_data: Dict[str, Any]):
        """Save enhanced processing job to Firestore"""
        if not FIREBASE_AVAILABLE or not firebase_service.is_available():
            print("⚠️ Firebase not available, skipping enhanced job save")
            return
            
        try:
            db = firebase_service.get_firestore_client()
            db.collection('enhanced_processing_jobs').document(job_data['job_id']).set(job_data)
            print(f"💾 Saved enhanced processing job {job_data['job_id']} to Firestore")
        except Exception as e:
            print(f"❌ Error saving enhanced processing job: {str(e)}")
    
    async def _save_enhanced_vector_chunks(self, job_id: str, user_id: str, chunks: List[EnhancedChunk]):
        """Save enhanced vector chunks to Firestore"""
        if not FIREBASE_AVAILABLE or not firebase_service.is_available():
            print("⚠️ Firebase not available, skipping enhanced vector chunk save")
            return
            
        try:
            db = firebase_service.get_firestore_client()
            
            for chunk in chunks:
                # Convert enhanced chunk to Firestore-compatible dict
                chunk_data = {
                    'id': chunk.id,
                    'job_id': job_id,
                    'user_id': user_id,
                    'text': chunk.text,
                    'embedding': chunk.embedding,
                    'embedding_model': chunk.embedding_model,
                    'prev_chunk_id': chunk.prev_chunk_id,
                    'next_chunk_id': chunk.next_chunk_id,
                    'related_chunks': chunk.related_chunks,
                    'metadata': {
                        'document_id': chunk.metadata.document_id,
                        'document_name': chunk.metadata.document_name,
                        'location': {
                            'page_number': chunk.metadata.location.page_number,
                            'bounding_box': chunk.metadata.location.bounding_box,
                            'paragraph_index': chunk.metadata.location.paragraph_index,
                            'line_range': chunk.metadata.location.line_range
                        },
                        'chunk_type': chunk.metadata.chunk_type.value,
                        'hierarchy_level': chunk.metadata.hierarchy_level,
                        'parent_section': chunk.metadata.parent_section,
                        'global_chunk_index': chunk.metadata.global_chunk_index,
                        'page_chunk_index': chunk.metadata.page_chunk_index,
                        'word_count': chunk.metadata.word_count,
                        'char_count': chunk.metadata.char_count,
                        'sentence_count': chunk.metadata.sentence_count,
                        'processing_version': chunk.metadata.processing_version,
                        'extracted_at': chunk.metadata.extracted_at
                    },
                    'created_at': datetime.now()
                }
                
                db.collection('enhanced_vector_chunks').document(chunk.id).set(chunk_data)
            
            print(f"💾 Saved {len(chunks)} enhanced vector chunks to Firestore")
            
        except Exception as e:
            print(f"❌ Error saving enhanced vector chunks: {str(e)}")
    
    async def get_enhanced_processing_jobs(self, user_id: str) -> List[Dict[str, Any]]:
        """Get enhanced processing jobs for a user"""
        if not FIREBASE_AVAILABLE or not firebase_service.is_available():
            print("⚠️ Firebase not available, returning empty enhanced jobs list")
            return []
            
        try:
            db = firebase_service.get_firestore_client()
            jobs_ref = db.collection('enhanced_processing_jobs')
            query = jobs_ref.where('user_id', '==', user_id).order_by('started_at', direction='DESCENDING').limit(20)
            
            jobs = []
            for doc in query.stream():
                job_data = doc.to_dict()
                # Convert Firestore timestamps to ISO strings if needed
                if 'started_at' in job_data and hasattr(job_data['started_at'], 'isoformat'):
                    job_data['started_at'] = job_data['started_at'].isoformat()
                if 'updated_at' in job_data and hasattr(job_data['updated_at'], 'isoformat'):
                    job_data['updated_at'] = job_data['updated_at'].isoformat()
                    
                jobs.append(job_data)
            
            print(f"📋 Retrieved {len(jobs)} enhanced processing jobs for user {user_id}")
            return jobs
            
        except Exception as e:
            print(f"❌ Error getting enhanced processing jobs: {str(e)}")
            return []
    
    async def vector_search(
        self, 
        user_id: str, 
        query: str, 
        filters: Optional[Dict[str, Any]] = None,
        include_metadata: bool = True
    ) -> List[Dict[str, Any]]:
        """Perform semantic vector search using embeddings"""
        print(f"🔍 Performing vector search")
        
        if not FIREBASE_AVAILABLE or not firebase_service.is_available():
            print("⚠️ Firebase not available, returning dummy results")
            return self._get_dummy_search_results(query)
        
        try:
            # Generate embedding for the search query
            query_embedding = await self._generate_query_embedding(query)
            
            if not query_embedding:
                print("⚠️ Could not generate query embedding, returning dummy results")
                return self._get_dummy_search_results(query)
            
            # Get vector chunks from Firestore
            db = firebase_service.get_firestore_client()
            chunks_ref = db.collection('vector_chunks').where('user_id', '==', user_id).limit(100)
            
            # Calculate similarities
            results = []
            for doc in chunks_ref.stream():
                chunk_data = doc.to_dict()
                
                if 'embedding' in chunk_data:
                    similarity = self._calculate_cosine_similarity(query_embedding, chunk_data['embedding'])
                    
                    # Apply similarity threshold filter
                    threshold = filters.get('similarity_threshold', 0.5) if filters else 0.5
                    if similarity >= threshold:
                        result = {
                            'content': chunk_data['text'],
                            'similarity': similarity,
                            'metadata': {
                                'job_id': chunk_data['job_id'],
                                'chunk_id': chunk_data['id'],
                                'word_count': chunk_data.get('metadata', {}).get('word_count', 0)
                            }
                        }
                        
                        if include_metadata:
                            result['metadata'].update(chunk_data.get('metadata', {}))
                        
                        results.append(result)
            
            # Sort by similarity and apply max_results filter
            results.sort(key=lambda x: x['similarity'], reverse=True)
            max_results = filters.get('max_results', 10) if filters else 10
            results = results[:max_results]
            
            print(f"✅ Found {len(results)} relevant chunks")
            return results
            
        except Exception as e:
            print(f"❌ Vector search error: {str(e)}")
            return self._get_dummy_search_results(query)
    
    def _get_dummy_search_results(self, query: str) -> List[Dict[str, Any]]:
        """Return dummy search results as fallback"""
        return [
            {
                'content': f"This is a sample search result for query: {query}",
                'similarity': 0.85,
                'metadata': {
                    'source': 'sample_document.pdf',
                    'page': 1
                }
            },
            {
                'content': f"Another relevant result matching: {query}",
                'similarity': 0.78,
                'metadata': {
                    'source': 'another_document.pdf', 
                    'page': 3
                }
            }
        ]
    
    async def _generate_query_embedding(self, query: str) -> Optional[List[float]]:
        """Generate embedding for search query"""
        try:
            if not self.embedding_model or not VERTEX_AI_AVAILABLE:
                print("⚠️ Vertex AI embeddings not available for query")
                return None
            
            embeddings = self.embedding_model.get_embeddings([query])
            return embeddings[0].values
            
        except Exception as e:
            print(f"❌ Query embedding error: {str(e)}")
            return None
    
    def _calculate_cosine_similarity(self, vec1: List[float], vec2: List[float]) -> float:
        """Calculate cosine similarity between two vectors"""
        try:
            import math
            
            # Calculate dot product
            dot_product = sum(a * b for a, b in zip(vec1, vec2))
            
            # Calculate magnitudes
            magnitude1 = math.sqrt(sum(a * a for a in vec1))
            magnitude2 = math.sqrt(sum(b * b for b in vec2))
            
            # Avoid division by zero
            if magnitude1 == 0 or magnitude2 == 0:
                return 0.0
            
            # Calculate cosine similarity
            similarity = dot_product / (magnitude1 * magnitude2)
            return max(0.0, min(1.0, similarity))  # Clamp between 0 and 1
            
        except Exception as e:
            print(f"❌ Similarity calculation error: {str(e)}")
            return 0.0
    
    async def get_knowledge_graph(
        self,
        user_id: str,
        entity_type: Optional[str] = None,
        search_query: Optional[str] = None,
        limit: int = 100
    ) -> Dict[str, List[Dict[str, Any]]]:
        """Get knowledge graph data from Firestore"""
        print(f"🕸️ Getting knowledge graph")
        
        if not FIREBASE_AVAILABLE or not firebase_service.is_available():
            print("⚠️ Firebase not available, returning dummy knowledge graph")
            return self._get_dummy_knowledge_graph()
        
        try:
            db = firebase_service.get_firestore_client()
            
            # Query entities
            entities_query = db.collection('knowledge_entities').where('user_id', '==', user_id)
            if entity_type:
                entities_query = entities_query.where('type', '==', entity_type)
            entities_query = entities_query.limit(limit)
            
            entities = []
            for doc in entities_query.stream():
                entity_data = doc.to_dict()
                entity_data['id'] = doc.id
                
                # Convert Firestore timestamp to ISO string
                if 'created_at' in entity_data and hasattr(entity_data['created_at'], 'isoformat'):
                    entity_data['created_at'] = entity_data['created_at'].isoformat()
                
                # Filter by search query if provided
                if search_query:
                    search_lower = search_query.lower()
                    if (search_lower not in entity_data.get('name', '').lower() and 
                        search_lower not in entity_data.get('description', '').lower()):
                        continue
                
                entities.append(entity_data)
            
            # Query relations
            relations_query = db.collection('knowledge_relations').where('user_id', '==', user_id).limit(limit * 2)
            relations = []
            for doc in relations_query.stream():
                relation_data = doc.to_dict()
                relation_data['id'] = doc.id
                
                # Convert Firestore timestamp to ISO string
                if 'created_at' in relation_data and hasattr(relation_data['created_at'], 'isoformat'):
                    relation_data['created_at'] = relation_data['created_at'].isoformat()
                
                relations.append(relation_data)
            
            print(f"✅ Retrieved {len(entities)} entities and {len(relations)} relations")
            return {
                'entities': entities,
                'relations': relations
            }
            
        except Exception as e:
            print(f"❌ Error getting knowledge graph: {str(e)}")
            return self._get_dummy_knowledge_graph()
    
    def _get_dummy_knowledge_graph(self) -> Dict[str, List[Dict[str, Any]]]:
        """Return dummy knowledge graph as fallback"""
        dummy_entities = [
            {
                'id': 'entity_1',
                'name': 'Machine Learning',
                'type': 'CONCEPT',
                'description': 'A method of data analysis that automates analytical model building',
                'confidence': 0.9
            },
            {
                'id': 'entity_2', 
                'name': 'Python Programming',
                'type': 'CONCEPT',
                'description': 'A high-level programming language',
                'confidence': 0.85
            }
        ]
        
        dummy_relations = [
            {
                'id': 'rel_1',
                'from_entity': 'entity_1',
                'to_entity': 'entity_2',
                'relation': 'implemented_using',
                'confidence': 0.8
            }
        ]
        
        return {
            'entities': dummy_entities,
            'relations': dummy_relations
        }
    
    async def get_processing_jobs(self, user_id: str) -> List[Dict[str, Any]]:
        """Get processing jobs for user from Firestore"""
        print(f"⚙️ Getting processing jobs")
        
        if not FIREBASE_AVAILABLE or not firebase_service.is_available():
            print("⚠️ Firebase not available, returning empty list")
            return []
            
        try:
            db = firebase_service.get_firestore_client()
            
            # Query processing jobs for user, ordered by creation time
            jobs_ref = db.collection('processing_jobs').where('user_id', '==', user_id).order_by('created_at', direction='DESCENDING').limit(20)
            
            jobs = []
            for doc in jobs_ref.stream():
                job_data = doc.to_dict()
                job_data['id'] = doc.id
                
                # Convert Firestore timestamp to ISO string
                if 'created_at' in job_data and hasattr(job_data['created_at'], 'isoformat'):
                    job_data['created_at'] = job_data['created_at'].isoformat()
                if 'completed_at' in job_data and hasattr(job_data['completed_at'], 'isoformat'):
                    job_data['completed_at'] = job_data['completed_at'].isoformat()
                    
                jobs.append(job_data)
            
            print(f"📊 Found {len(jobs)} processing jobs")
            return jobs
            
        except Exception as e:
            print(f"❌ Error getting processing jobs: {str(e)}")
            return []
    
    async def get_entity_details(self, user_id: str, entity_id: str) -> Optional[Dict[str, Any]]:
        """Get entity details - returns dummy data for now"""
        print(f"📋 Getting entity details for {entity_id}")
        
        return {
            'id': entity_id,
            'name': 'Sample Entity',
            'type': 'CONCEPT',
            'description': 'This is a sample entity for demonstration',
            'attributes': {
                'category': 'Technology',
                'importance': 'High'
            },
            'confidence': 0.85
        }
    
    async def get_processed_documents(self, user_id: str) -> List[Dict[str, Any]]:
        """Get processed documents for user from Firestore"""
        print(f"📚 Getting processed documents")
        
        if not FIREBASE_AVAILABLE or not firebase_service.is_available():
            print("⚠️ Firebase not available, returning empty list")
            return []
            
        try:
            db = firebase_service.get_firestore_client()
            
            # Query processed documents for user
            docs_ref = db.collection('processed_documents').where('user_id', '==', user_id).order_by('processed_at', direction='DESCENDING').limit(50)
            
            documents = []
            for doc in docs_ref.stream():
                doc_data = doc.to_dict()
                doc_data['id'] = doc.id
                
                # Convert Firestore timestamp to ISO string
                if 'processed_at' in doc_data and hasattr(doc_data['processed_at'], 'isoformat'):
                    doc_data['processed_at'] = doc_data['processed_at'].isoformat()
                    
                documents.append(doc_data)
            
            print(f"📊 Found {len(documents)} processed documents")
            return documents
            
        except Exception as e:
            print(f"❌ Error getting processed documents: {str(e)}")
            return []
    
    async def delete_document(self, user_id: str, document_id: str) -> bool:
        """Delete document - returns True for now"""
        print(f"🗑️ Deleting document {document_id}")
        return True
    
    async def get_knowledge_stats(self, user_id: str) -> Dict[str, Any]:
        """Get knowledge base statistics from Firestore"""
        print(f"📊 Getting knowledge stats")
        
        if not FIREBASE_AVAILABLE or not firebase_service.is_available():
            print("⚠️ Firebase not available, returning empty stats")
            return {
                'total_documents': 0,
                'total_chunks': 0, 
                'total_entities': 0,
                'total_relations': 0
            }
            
        try:
            db = firebase_service.get_firestore_client()
            
            # Count documents
            docs_count = len([doc for doc in db.collection('processed_documents').where('user_id', '==', user_id).stream()])
            
            # Count vector chunks
            chunks_count = len([doc for doc in db.collection('vector_chunks').where('user_id', '==', user_id).stream()])
            
            # Count entities
            entities_count = len([doc for doc in db.collection('knowledge_entities').where('user_id', '==', user_id).stream()])
            
            # Count relations
            relations_count = len([doc for doc in db.collection('knowledge_relations').where('user_id', '==', user_id).stream()])
            
            stats = {
                'total_documents': docs_count,
                'total_chunks': chunks_count,
                'total_entities': entities_count,
                'total_relations': relations_count
            }
            
            print(f"📊 Knowledge stats retrieved successfully")
            return stats
            
        except Exception as e:
            print(f"❌ Error getting knowledge stats: {str(e)}")
            return {
                'total_documents': 0,
                'total_chunks': 0, 
                'total_entities': 0,
                'total_relations': 0
            }
    
    async def reprocess_document(
        self,
        job_id: str,
        user_id: str, 
        document_id: str,
        options: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Reprocess document - returns success for now"""
        print(f"🔄 Reprocessing document {document_id}")
        
        return {
            'id': job_id,
            'user_id': user_id,
            'document_id': document_id,
            'status': 'completed',
            'message': 'Document reprocessed successfully'
        }

    # Firestore helper methods
    async def _save_processing_job(self, job_data: Dict[str, Any]):
        """Save processing job to Firestore"""
        if not FIREBASE_AVAILABLE or not firebase_service.is_available():
            print("⚠️ Firebase not available, skipping job save")
            return
            
        try:
            db = firebase_service.get_firestore_client()
            db.collection('processing_jobs').document(job_data['id']).set(job_data)
            print(f"💾 Saved processing job {job_data['id']} to Firestore")
        except Exception as e:
            print(f"❌ Error saving processing job: {str(e)}")
    
    async def _update_processing_job(self, job_id: str, updates: Dict[str, Any]):
        """Update processing job in Firestore"""
        if not FIREBASE_AVAILABLE or not firebase_service.is_available():
            print("⚠️ Firebase not available, skipping job update")
            return
            
        try:
            db = firebase_service.get_firestore_client()
            db.collection('processing_jobs').document(job_id).update(updates)
            print(f"🔄 Updated processing job {job_id} in Firestore")
        except Exception as e:
            print(f"❌ Error updating processing job: {str(e)}")
    
    async def _update_job_progress(self, job_id: str, progress: int, step: str):
        """Update job progress"""
        updates = {
            'progress': progress,
            'current_step': step,
            'updated_at': datetime.now()
        }
        await self._update_processing_job(job_id, updates)
        print(f"📈 Job {job_id} progress: {progress}% - {step}")
    
    async def _save_processed_document(self, doc_data: Dict[str, Any]):
        """Save processed document to Firestore"""
        if not FIREBASE_AVAILABLE or not firebase_service.is_available():
            print("⚠️ Firebase not available, skipping document save")
            return
            
        try:
            db = firebase_service.get_firestore_client()
            db.collection('processed_documents').document(doc_data['id']).set(doc_data)
            print(f"📄 Saved processed document {doc_data['id']} to Firestore")
        except Exception as e:
            print(f"❌ Error saving processed document: {str(e)}")
    
    async def _upload_to_storage(self, job_id: str, user_id: str, content: bytes, file_name: str, content_type: str) -> str:
        """Upload file to Cloud Storage or save locally"""
        
        if self.storage_client and STORAGE_AVAILABLE:
            try:
                # Use a bucket name based on project
                project_id = settings.google_cloud_project or settings.firebase_project_id
                bucket_name = f"{project_id}-knowledge-docs"
                
                # Get or create bucket
                try:
                    bucket = self.storage_client.bucket(bucket_name)
                    if not bucket.exists():
                        bucket = self.storage_client.create_bucket(bucket_name, location="us-central1")
                except:
                    bucket = self.storage_client.bucket(bucket_name)
                
                # Upload file
                blob_name = f"users/{user_id}/documents/{job_id}/{file_name}"
                blob = bucket.blob(blob_name)
                blob.upload_from_string(content, content_type=content_type)
                
                storage_uri = f"gs://{bucket_name}/{blob_name}"
                print(f"✅ Uploaded to Cloud Storage: {storage_uri}")
                return storage_uri
                
            except Exception as e:
                print(f"⚠️ Cloud Storage upload failed: {str(e)}, falling back to local storage")
        
        # Fallback: save to local temp file
        import tempfile
        temp_dir = tempfile.mkdtemp()
        temp_path = os.path.join(temp_dir, f"{job_id}_{file_name}")
        
        with open(temp_path, 'wb') as f:
            f.write(content)
        
        print(f"✅ Saved locally: {temp_path}")
        return temp_path
    
    async def _extract_text_with_docai(self, file_path: str, content_type: str, filename: str) -> str:
        """Extract text using Document AI or fallback method"""
        
        # Check if Document AI is available and configured
        project_id = settings.google_cloud_project or settings.firebase_project_id
        location = settings.vertex_ai_location
        
        if (self.doc_ai_client and DOCUMENT_AI_AVAILABLE and 
            project_id and location):
            
            try:
                return await self._extract_with_docai(file_path, content_type)
            except Exception as e:
                print(f"⚠️ Document AI failed: {str(e)}, using fallback")
        
        # Fallback to simple text extraction
        return await self._extract_text_fallback(file_path, content_type, filename)
    
    async def _extract_with_docai(self, file_path: str, content_type: str) -> str:
        """Extract text using Document AI"""
        # For Document AI, we need a processor ID
        # This would be configured in production
        project_id = settings.google_cloud_project or settings.firebase_project_id
        location = settings.vertex_ai_location
        processor_name = f"projects/{project_id}/locations/{location}/processors/YOUR_PROCESSOR_ID"
        
        # Read file content
        if file_path.startswith('gs://'):
            # File is in Cloud Storage
            bucket_name = file_path.split('/')[2]
            blob_name = '/'.join(file_path.split('/')[3:])
            bucket = self.storage_client.bucket(bucket_name)
            blob = bucket.blob(blob_name)
            content = blob.download_as_bytes()
        else:
            # Local file
            with open(file_path, 'rb') as f:
                content = f.read()
        
        # Create Document AI request
        raw_document = documentai.RawDocument(
            content=content,
            mime_type=content_type
        )
        
        request = documentai.ProcessRequest(
            name=processor_name,
            raw_document=raw_document
        )
        
        # Process document
        result = self.doc_ai_client.process_document(request=request)
        document = result.document
        
        print(f"✅ Document AI extracted {len(document.text)} characters")
        return document.text
    
    async def _extract_text_fallback(self, file_path: str, content_type: str, filename: str) -> str:
        """Fallback text extraction using basic methods"""
        try:
            if content_type == 'application/pdf':
                # Try PyPDF2 for PDFs
                try:
                    import PyPDF2
                    with open(file_path, 'rb') as file:
                        reader = PyPDF2.PdfReader(file)
                        text = ""
                        for page in reader.pages:
                            text += page.extract_text() + "\n"
                        print(f"✅ Extracted {len(text)} characters from PDF using PyPDF2")
                        return text
                except ImportError:
                    print("⚠️ PyPDF2 not available")
                except Exception as e:
                    print(f"⚠️ PyPDF2 extraction failed: {str(e)}")
            
            elif content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                # Try python-docx for DOCX
                try:
                    import docx
                    doc = docx.Document(file_path)
                    text = ""
                    for paragraph in doc.paragraphs:
                        text += paragraph.text + "\n"
                    print(f"✅ Extracted {len(text)} characters from DOCX using python-docx")
                    return text
                except ImportError:
                    print("⚠️ python-docx not available")
                except Exception as e:
                    print(f"⚠️ python-docx extraction failed: {str(e)}")
            
            # Final fallback - return basic info
            fallback_text = f"Document: {filename}\nContent type: {content_type}\nProcessed at: {datetime.now()}\n\nText extraction not available. Please install appropriate libraries or configure Document AI."
            print(f"⚠️ Using fallback text for {filename}")
            return fallback_text
            
        except Exception as e:
            print(f"❌ Text extraction error: {str(e)}")
            return f"Error extracting text from {filename}: {str(e)}"
    
    async def _chunk_text(self, text: str, chunk_size: int, overlap_size: int) -> List[Dict[str, Any]]:
        """Split text into overlapping chunks"""
        try:
            # Split by sentences first, then by words if needed
            sentences = text.split('. ')
            chunks = []
            current_chunk = ""
            chunk_id = 0
            
            for sentence in sentences:
                # Check if adding this sentence would exceed chunk size
                test_chunk = current_chunk + ". " + sentence if current_chunk else sentence
                
                if len(test_chunk.split()) <= chunk_size:
                    current_chunk = test_chunk
                else:
                    # Save current chunk if it's not empty
                    if current_chunk.strip():
                        chunks.append({
                            'id': f"{chunk_id}",
                            'text': current_chunk.strip(),
                            'word_count': len(current_chunk.split()),
                            'start_index': chunk_id * (chunk_size - overlap_size),
                            'end_index': chunk_id * (chunk_size - overlap_size) + len(current_chunk.split())
                        })
                        chunk_id += 1
                    
                    # Start new chunk
                    current_chunk = sentence
            
            # Add the last chunk if not empty
            if current_chunk.strip():
                chunks.append({
                    'id': f"{chunk_id}",
                    'text': current_chunk.strip(),
                    'word_count': len(current_chunk.split()),
                    'start_index': chunk_id * (chunk_size - overlap_size),
                    'end_index': chunk_id * (chunk_size - overlap_size) + len(current_chunk.split())
                })
            
            print(f"✅ Created {len(chunks)} text chunks")
            return chunks
            
        except Exception as e:
            print(f"❌ Text chunking error: {str(e)}")
            # Return single chunk as fallback
            return [{
                'id': '0',
                'text': text[:5000],  # Limit to 5000 chars
                'word_count': len(text.split()),
                'start_index': 0,
                'end_index': len(text.split())
            }]
    
    async def _generate_embeddings(self, chunks: List[Dict[str, Any]]) -> List[List[float]]:
        """Generate embeddings using Vertex AI"""
        try:
            if not self.embedding_model or not VERTEX_AI_AVAILABLE:
                print("⚠️ Vertex AI embeddings not available, using dummy embeddings")
                # Return dummy embeddings with consistent dimensions (768 for text-multilingual-embedding-002)
                return [[0.1] * 768 for _ in chunks]
            
            print(f"🔢 Generating embeddings for {len(chunks)} chunks using Vertex AI")
            
            # Extract text from chunks
            texts = [chunk['text'] for chunk in chunks]
            
            # Generate embeddings in batches to avoid API limits
            batch_size = 5
            all_embeddings = []
            
            for i in range(0, len(texts), batch_size):
                batch_texts = texts[i:i + batch_size]
                
                try:
                    embeddings = self.embedding_model.get_embeddings(batch_texts)
                    batch_embeddings = [embedding.values for embedding in embeddings]
                    all_embeddings.extend(batch_embeddings)
                    
                    print(f"✅ Generated embeddings for batch {i//batch_size + 1}/{(len(texts)-1)//batch_size + 1}")
                    
                except Exception as e:
                    print(f"⚠️ Batch embedding failed: {str(e)}, using dummy embeddings for batch")
                    # Use dummy embeddings for failed batch
                    dummy_batch = [[0.1] * 768 for _ in batch_texts]
                    all_embeddings.extend(dummy_batch)
            
            print(f"✅ Generated {len(all_embeddings)} embeddings")
            return all_embeddings
            
        except Exception as e:
            print(f"❌ Embedding generation error: {str(e)}")
            # Return dummy embeddings as fallback
            return [[0.1] * 768 for _ in chunks]
    
    async def _save_vector_chunks(self, job_id: str, user_id: str, chunks: List[Dict[str, Any]], embeddings: List[List[float]]):
        """Save vector chunks to Firestore"""
        if not FIREBASE_AVAILABLE or not firebase_service.is_available():
            print("⚠️ Firebase not available, skipping vector chunk save")
            return
            
        try:
            db = firebase_service.get_firestore_client()
            
            for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
                chunk_data = {
                    'id': f"{job_id}_{chunk['id']}",
                    'job_id': job_id,
                    'user_id': user_id,
                    'text': chunk['text'],
                    'embedding': embedding,
                    'metadata': {
                        'word_count': chunk['word_count'],
                        'start_index': chunk['start_index'],
                        'end_index': chunk['end_index'],
                        'chunk_index': i
                    },
                    'created_at': datetime.now()
                }
                
                db.collection('vector_chunks').document(chunk_data['id']).set(chunk_data)
            
            print(f"💾 Saved {len(chunks)} vector chunks to Firestore")
            
        except Exception as e:
            print(f"❌ Error saving vector chunks: {str(e)}")
    
    async def _extract_knowledge_graph(self, text: str, chunks: List[Dict[str, Any]]) -> Dict[str, List[Dict[str, Any]]]:
        """Extract knowledge graph using Gemini API"""
        try:
            if not GEMINI_AVAILABLE or not gemini_service:
                print("⚠️ Gemini service not available, using dummy knowledge graph")
                return self._get_dummy_knowledge_extraction()
            
            print("🧠 Extracting knowledge graph using Gemini API")
            
            # Prepare text for analysis (use first few chunks to avoid token limits)
            analysis_text = text[:8000]  # Limit to ~8k chars to stay within token limits
            
            # Create knowledge graph extraction prompt
            kg_prompt = f"""
あなたは高度な知識グラフ抽出システムです。以下のテキストから構造化された知識グラフを抽出してください。

テキスト:
{analysis_text}

以下の形式でJSONを返してください:
{{
    "entities": [
        {{
            "name": "エンティティ名",
            "type": "PERSON|ORGANIZATION|CONCEPT|LOCATION|EVENT|TECHNOLOGY|OTHER",
            "description": "エンティティの説明",
            "confidence": 0.0-1.0の信頼度
        }}
    ],
    "relations": [
        {{
            "from_entity": "エンティティ1の名前",
            "to_entity": "エンティティ2の名前", 
            "relation": "関係性の種類",
            "confidence": 0.0-1.0の信頼度
        }}
    ]
}}

重要な要求:
1. 有効なJSONのみを返してください
2. 最も重要で明確なエンティティと関係性のみを抽出してください
3. 信頼度は実際の確信度を反映してください
4. エンティティ名は正規化してください（略語を展開、表記統一）
5. 最大20個のエンティティと30個の関係性まで
"""
            
            # Send request to Gemini
            response = await gemini_service.send_message(
                model_name="gemini-2-0-flash-001",
                history=[],
                message=kg_prompt
            )
            
            # Parse JSON response
            try:
                import json
                import re
                
                # Extract JSON from response
                response_text = response  # send_message returns string directly
                
                # Find JSON in response (handle markdown code blocks)
                json_match = re.search(r'```(?:json)?\s*(\{.*?\})\s*```', response_text, re.DOTALL)
                if json_match:
                    json_text = json_match.group(1)
                else:
                    # Try to find JSON without code blocks
                    json_match = re.search(r'(\{.*\})', response_text, re.DOTALL)
                    if json_match:
                        json_text = json_match.group(1)
                    else:
                        raise ValueError("No JSON found in response")
                
                knowledge_graph = json.loads(json_text)
                
                # Validate structure
                if 'entities' not in knowledge_graph or 'relations' not in knowledge_graph:
                    raise ValueError("Invalid knowledge graph structure")
                
                # Add IDs to entities and relations
                for i, entity in enumerate(knowledge_graph['entities']):
                    entity['id'] = f"entity_{i}"
                
                for i, relation in enumerate(knowledge_graph['relations']):
                    relation['id'] = f"relation_{i}"
                
                print(f"✅ Extracted {len(knowledge_graph['entities'])} entities and {len(knowledge_graph['relations'])} relations")
                return knowledge_graph
                
            except (json.JSONDecodeError, ValueError) as e:
                print(f"❌ Failed to parse knowledge graph JSON: {str(e)}")
                return self._get_dummy_knowledge_extraction()
            
        except Exception as e:
            print(f"❌ Knowledge graph extraction error: {str(e)}")
            return self._get_dummy_knowledge_extraction()
    
    def _get_dummy_knowledge_extraction(self) -> Dict[str, List[Dict[str, Any]]]:
        """Return dummy knowledge graph extraction"""
        return {
            "entities": [
                {
                    "id": "entity_0",
                    "name": "サンプルコンセプト",
                    "type": "CONCEPT",
                    "description": "テキストから抽出されたサンプルエンティティ",
                    "confidence": 0.8
                }
            ],
            "relations": []
        }
    
    async def _save_knowledge_graph(self, job_id: str, user_id: str, entities: List[Dict[str, Any]], relations: List[Dict[str, Any]]):
        """Save knowledge graph entities and relations to Firestore"""
        if not FIREBASE_AVAILABLE or not firebase_service.is_available():
            print("⚠️ Firebase not available, skipping knowledge graph save")
            return
            
        try:
            db = firebase_service.get_firestore_client()
            
            # Save entities
            for entity in entities:
                entity_data = {
                    'id': f"{job_id}_{entity['id']}",
                    'job_id': job_id,
                    'user_id': user_id,
                    'name': entity['name'],
                    'type': entity['type'],
                    'description': entity.get('description', ''),
                    'confidence': entity.get('confidence', 0.0),
                    'created_at': datetime.now()
                }
                
                db.collection('knowledge_entities').document(entity_data['id']).set(entity_data)
            
            # Save relations
            for relation in relations:
                relation_data = {
                    'id': f"{job_id}_{relation['id']}",
                    'job_id': job_id,
                    'user_id': user_id,
                    'from_entity': relation['from_entity'],
                    'to_entity': relation['to_entity'],
                    'relation': relation['relation'],
                    'confidence': relation.get('confidence', 0.0),
                    'created_at': datetime.now()
                }
                
                db.collection('knowledge_relations').document(relation_data['id']).set(relation_data)
            
            print(f"💾 Saved {len(entities)} entities and {len(relations)} relations to Firestore")
            
        except Exception as e:
            print(f"❌ Error saving knowledge graph: {str(e)}")

# Create global instance
knowledge_service = KnowledgeService()