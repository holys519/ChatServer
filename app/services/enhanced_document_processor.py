"""
Enhanced Document Processing Service
Provides advanced text extraction with page-level location tracking
"""

import os
import re
from typing import List, Dict, Any, Optional, Tuple
from datetime import datetime

from app.models.enhanced_schemas import (
    EnhancedChunk, ChunkMetadata, DocumentLocation, ChunkType, 
    ProcessingStatus
)

# Document AI imports
try:
    from google.cloud import documentai
    DOCUMENT_AI_AVAILABLE = True
except ImportError:
    documentai = None
    DOCUMENT_AI_AVAILABLE = False

# Fallback imports
try:
    import PyPDF2
    import fitz  # PyMuPDF for better PDF processing
    PDF_AVAILABLE = True
except ImportError:
    PyPDF2 = None
    fitz = None
    PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DocxDocument = None
    DOCX_AVAILABLE = False

class EnhancedDocumentProcessor:
    """Enhanced document processor with location tracking"""
    
    def __init__(self, project_id: Optional[str] = None, location: str = "us-central1"):
        self.project_id = project_id
        self.location = location
        self.client = None
        
        # Initialize Document AI if available
        if DOCUMENT_AI_AVAILABLE and project_id:
            try:
                self.client = documentai.DocumentProcessorServiceClient()
                print("✅ Document AI client initialized")
            except Exception as e:
                print(f"⚠️ Document AI initialization failed: {e}")
                self.client = None
    
    async def process_document_with_location(
        self, 
        file_content: bytes, 
        file_name: str, 
        content_type: str,
        job_id: str,
        chunk_size: int = 500,
        overlap_size: int = 50
    ) -> Tuple[List[EnhancedChunk], ProcessingStatus]:
        """
        Process document with enhanced location tracking
        
        Returns:
            Tuple of (enhanced_chunks, processing_status)
        """
        print(f"🔍 Processing document with enhanced location tracking: {file_name}")
        
        # Initialize processing status
        status = ProcessingStatus(
            job_id=job_id,
            user_id="", # Will be set by caller
            document_name=file_name,
            status="processing",
            progress=10,
            current_step="Analyzing document structure",
            started_at=datetime.now(),
            updated_at=datetime.now()
        )
        
        try:
            # Determine processing method
            if content_type == "application/pdf":
                chunks, pages_info = await self._process_pdf_with_locations(
                    file_content, file_name, job_id, chunk_size, overlap_size
                )
            elif "wordprocessingml" in content_type:
                chunks, pages_info = await self._process_docx_with_locations(
                    file_content, file_name, job_id, chunk_size, overlap_size
                )
            else:
                raise ValueError(f"Unsupported document type: {content_type}")
            
            # Update status
            status.total_pages = pages_info.get("total_pages", 1)
            status.pages_processed = status.total_pages
            status.chunks_created = len(chunks)
            status.progress = 90
            status.current_step = "Finalizing chunk processing"
            status.updated_at = datetime.now()
            
            # Link chunks for context (previous/next relationships)
            self._link_chunks_for_context(chunks)
            
            status.progress = 100
            status.status = "completed"
            status.current_step = "Document processing completed"
            status.updated_at = datetime.now()
            
            print(f"✅ Document processed: {len(chunks)} chunks across {status.total_pages} pages")
            return chunks, status
            
        except Exception as e:
            status.status = "failed"
            status.error_message = str(e)
            status.updated_at = datetime.now()
            print(f"❌ Document processing failed: {e}")
            return [], status
    
    async def _process_pdf_with_locations(
        self, 
        file_content: bytes, 
        file_name: str, 
        job_id: str,
        chunk_size: int,
        overlap_size: int
    ) -> Tuple[List[EnhancedChunk], Dict[str, Any]]:
        """Process PDF with page and coordinate information"""
        
        # Try Document AI first
        if self.client and DOCUMENT_AI_AVAILABLE:
            try:
                return await self._process_pdf_with_document_ai(
                    file_content, file_name, job_id, chunk_size, overlap_size
                )
            except Exception as e:
                print(f"⚠️ Document AI failed, falling back to PyMuPDF: {e}")
        
        # Fallback to PyMuPDF for better PDF processing
        if fitz:
            return await self._process_pdf_with_pymupdf(
                file_content, file_name, job_id, chunk_size, overlap_size
            )
        
        # Final fallback to PyPDF2
        if PDF_AVAILABLE:
            return await self._process_pdf_with_pypdf2(
                file_content, file_name, job_id, chunk_size, overlap_size
            )
        
        raise ValueError("No PDF processing library available")
    
    async def _process_pdf_with_document_ai(
        self, 
        file_content: bytes, 
        file_name: str, 
        job_id: str,
        chunk_size: int,
        overlap_size: int
    ) -> Tuple[List[EnhancedChunk], Dict[str, Any]]:
        """Process PDF using Google Document AI for precise extraction"""
        
        # This would require a Document AI processor setup
        # For now, we'll implement a placeholder that falls back to PyMuPDF
        print("🚧 Document AI processing - implementation pending")
        return await self._process_pdf_with_pymupdf(
            file_content, file_name, job_id, chunk_size, overlap_size
        )
    
    async def _process_pdf_with_pymupdf(
        self, 
        file_content: bytes, 
        file_name: str, 
        job_id: str,
        chunk_size: int,
        overlap_size: int
    ) -> Tuple[List[EnhancedChunk], Dict[str, Any]]:
        """Process PDF using PyMuPDF with location tracking"""
        
        if not fitz:
            raise ImportError("PyMuPDF (fitz) not available")
        
        chunks = []
        pages_info = {"total_pages": 0}
        
        # Open PDF from bytes
        pdf_doc = fitz.open(stream=file_content, filetype="pdf")
        pages_info["total_pages"] = pdf_doc.page_count
        
        global_chunk_index = 0
        
        for page_num in range(pdf_doc.page_count):
            page = pdf_doc[page_num]
            page_chunk_index = 0
            
            # Extract text with location information
            blocks = page.get_text("dict")["blocks"]
            
            current_paragraph = ""
            current_bbox = None
            
            for block in blocks:
                if "lines" in block:  # Text block
                    for line in block["lines"]:
                        line_text = ""
                        line_bbox = line["bbox"]
                        
                        for span in line["spans"]:
                            line_text += span["text"]
                        
                        line_text = line_text.strip()
                        if line_text:
                            # Check if this starts a new paragraph
                            if self._is_paragraph_break(line_text, current_paragraph):
                                # Process current paragraph if it exists
                                if current_paragraph.strip():
                                    paragraph_chunks = self._split_text_into_chunks(
                                        current_paragraph.strip(), 
                                        chunk_size, 
                                        overlap_size
                                    )
                                    
                                    for i, chunk_text in enumerate(paragraph_chunks):
                                        chunk = self._create_enhanced_chunk(
                                            chunk_text=chunk_text,
                                            job_id=job_id,
                                            document_name=file_name,
                                            page_number=page_num + 1,
                                            global_index=global_chunk_index,
                                            page_index=page_chunk_index,
                                            bounding_box=current_bbox,
                                            chunk_type=ChunkType.PARAGRAPH
                                        )
                                        chunks.append(chunk)
                                        global_chunk_index += 1
                                        page_chunk_index += 1
                                
                                # Start new paragraph
                                current_paragraph = line_text
                                current_bbox = line_bbox
                            else:
                                # Continue current paragraph
                                current_paragraph += " " + line_text
                                if current_bbox:
                                    # Expand bounding box
                                    current_bbox = [
                                        min(current_bbox[0], line_bbox[0]),  # x0
                                        min(current_bbox[1], line_bbox[1]),  # y0
                                        max(current_bbox[2], line_bbox[2]),  # x1
                                        max(current_bbox[3], line_bbox[3])   # y1
                                    ]
                                else:
                                    current_bbox = line_bbox
            
            # Process final paragraph on page
            if current_paragraph.strip():
                paragraph_chunks = self._split_text_into_chunks(
                    current_paragraph.strip(), 
                    chunk_size, 
                    overlap_size
                )
                
                for chunk_text in paragraph_chunks:
                    chunk = self._create_enhanced_chunk(
                        chunk_text=chunk_text,
                        job_id=job_id,
                        document_name=file_name,
                        page_number=page_num + 1,
                        global_index=global_chunk_index,
                        page_index=page_chunk_index,
                        bounding_box=current_bbox,
                        chunk_type=ChunkType.PARAGRAPH
                    )
                    chunks.append(chunk)
                    global_chunk_index += 1
                    page_chunk_index += 1
        
        pdf_doc.close()
        return chunks, pages_info
    
    async def _process_pdf_with_pypdf2(
        self, 
        file_content: bytes, 
        file_name: str, 
        job_id: str,
        chunk_size: int,
        overlap_size: int
    ) -> Tuple[List[EnhancedChunk], Dict[str, Any]]:
        """Fallback PDF processing with PyPDF2 (limited location info)"""
        
        import io
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        chunks = []
        pages_info = {"total_pages": len(pdf_reader.pages)}
        
        global_chunk_index = 0
        
        for page_num, page in enumerate(pdf_reader.pages):
            text = page.extract_text()
            page_chunks = self._split_text_into_chunks(text, chunk_size, overlap_size)
            
            for page_chunk_index, chunk_text in enumerate(page_chunks):
                chunk = self._create_enhanced_chunk(
                    chunk_text=chunk_text,
                    job_id=job_id,
                    document_name=file_name,
                    page_number=page_num + 1,
                    global_index=global_chunk_index,
                    page_index=page_chunk_index,
                    bounding_box=None,  # PyPDF2 doesn't provide location info
                    chunk_type=ChunkType.PARAGRAPH
                )
                chunks.append(chunk)
                global_chunk_index += 1
        
        return chunks, pages_info
    
    async def _process_docx_with_locations(
        self, 
        file_content: bytes, 
        file_name: str, 
        job_id: str,
        chunk_size: int,
        overlap_size: int
    ) -> Tuple[List[EnhancedChunk], Dict[str, Any]]:
        """Process DOCX with paragraph and structure information"""
        
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not available")
        
        import io
        doc = DocxDocument(io.BytesIO(file_content))
        chunks = []
        
        # Estimate pages (rough calculation)
        total_chars = sum(len(para.text) for para in doc.paragraphs)
        estimated_pages = max(1, total_chars // 2000)  # Rough estimate
        pages_info = {"total_pages": estimated_pages}
        
        global_chunk_index = 0
        current_page = 1
        chars_on_page = 0
        page_chunk_index = 0
        
        for para_index, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue
            
            # Estimate page breaks (very rough)
            chars_on_page += len(text)
            if chars_on_page > 2000:  # Rough page break estimation
                current_page += 1
                chars_on_page = 0
                page_chunk_index = 0
            
            # Determine chunk type based on paragraph style
            chunk_type = ChunkType.PARAGRAPH
            hierarchy_level = None
            
            if paragraph.style.name.startswith('Heading'):
                chunk_type = ChunkType.HEADING
                try:
                    hierarchy_level = int(paragraph.style.name.split()[-1])
                except:
                    hierarchy_level = 1
            
            # Split paragraph into chunks if needed
            paragraph_chunks = self._split_text_into_chunks(text, chunk_size, overlap_size)
            
            for chunk_text in paragraph_chunks:
                chunk = self._create_enhanced_chunk(
                    chunk_text=chunk_text,
                    job_id=job_id,
                    document_name=file_name,
                    page_number=current_page,
                    global_index=global_chunk_index,
                    page_index=page_chunk_index,
                    bounding_box=None,  # DOCX doesn't easily provide coordinates
                    chunk_type=chunk_type,
                    hierarchy_level=hierarchy_level,
                    paragraph_index=para_index
                )
                chunks.append(chunk)
                global_chunk_index += 1
                page_chunk_index += 1
        
        return chunks, pages_info
    
    def _create_enhanced_chunk(
        self,
        chunk_text: str,
        job_id: str,
        document_name: str,
        page_number: int,
        global_index: int,
        page_index: int,
        bounding_box: Optional[List[float]] = None,
        chunk_type: ChunkType = ChunkType.PARAGRAPH,
        hierarchy_level: Optional[int] = None,
        paragraph_index: Optional[int] = None
    ) -> EnhancedChunk:
        """Create an enhanced chunk with full metadata"""
        
        # Convert bounding box to dict format if provided
        bbox_dict = None
        if bounding_box:
            bbox_dict = {
                "x": bounding_box[0],
                "y": bounding_box[1], 
                "width": bounding_box[2] - bounding_box[0],
                "height": bounding_box[3] - bounding_box[1]
            }
        
        # Create location info
        location = DocumentLocation(
            page_number=page_number,
            bounding_box=bbox_dict,
            paragraph_index=paragraph_index
        )
        
        # Calculate content statistics
        word_count = len(chunk_text.split())
        char_count = len(chunk_text)
        sentence_count = len([s for s in chunk_text.split('.') if s.strip()])
        
        # Create metadata
        metadata = ChunkMetadata(
            document_id=job_id,
            document_name=document_name,
            location=location,
            chunk_type=chunk_type,
            hierarchy_level=hierarchy_level,
            global_chunk_index=global_index,
            page_chunk_index=page_index,
            word_count=word_count,
            char_count=char_count,
            sentence_count=sentence_count
        )
        
        # Create chunk
        chunk_id = f"{job_id}_chunk_{global_index:04d}"
        
        return EnhancedChunk(
            id=chunk_id,
            text=chunk_text,
            metadata=metadata
        )
    
    def _split_text_into_chunks(
        self, 
        text: str, 
        chunk_size: int, 
        overlap_size: int
    ) -> List[str]:
        """Split text into overlapping chunks"""
        
        words = text.split()
        if len(words) <= chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(words):
            end = min(start + chunk_size, len(words))
            chunk_words = words[start:end]
            chunks.append(" ".join(chunk_words))
            
            if end >= len(words):
                break
            
            start = end - overlap_size
        
        return chunks
    
    def _link_chunks_for_context(self, chunks: List[EnhancedChunk]) -> None:
        """Link chunks to establish previous/next relationships"""
        
        for i, chunk in enumerate(chunks):
            if i > 0:
                chunk.prev_chunk_id = chunks[i-1].id
            if i < len(chunks) - 1:
                chunk.next_chunk_id = chunks[i+1].id
    
    def _is_paragraph_break(self, current_line: str, existing_paragraph: str) -> bool:
        """Determine if current line starts a new paragraph"""
        
        # Empty existing paragraph always starts new
        if not existing_paragraph.strip():
            return True
        
        # Lines that likely start new paragraphs
        paragraph_starters = [
            # Numbered lists
            r'^\d+\.',
            # Bullet points
            r'^[•·▪▫◦‣⁃]',
            # Letter lists
            r'^[a-zA-Z]\.',
            # Roman numerals
            r'^[ivxlcdm]+\.',
            # Headings (all caps or title case with limited words)
            r'^[A-Z][A-Z\s]{2,20}$',
            # New sentence after period + significant whitespace
            r'^[A-Z]'
        ]
        
        # Check if current line matches paragraph starter patterns
        for pattern in paragraph_starters:
            if re.match(pattern, current_line, re.IGNORECASE):
                return True
        
        # Check if previous paragraph ended with sentence-ending punctuation
        if existing_paragraph.rstrip().endswith(('.', '!', '?', ':')):
            return True
        
        return False